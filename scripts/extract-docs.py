#!/usr/bin/env python3
"""Extract text from PDF and DOCX files in docs/ for planning."""

from __future__ import annotations

import json
import re
import sys
from datetime import datetime, timezone
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = ROOT / "docs"
OUTPUT_DIR = ROOT / "docs" / "extracted"


def extract_pdf(path: Path) -> dict:
    try:
        import pdfplumber
    except ImportError:
        print("Installing pdfplumber...", file=sys.stderr)
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "pdfplumber", "-q"])
        import pdfplumber

    pages: list[dict] = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            tables = page.extract_tables() or []
            pages.append({
                "page": i,
                "text": text.strip(),
                "tables": [
                    [[cell or "" for cell in row] for row in table]
                    for table in tables
                    if table
                ],
            })

    full_text = "\n\n".join(
        f"--- Page {p['page']} ---\n{p['text']}" for p in pages if p["text"]
    )
    return {
        "source": path.name,
        "type": "pdf",
        "page_count": len(pages),
        "pages": pages,
        "full_text": full_text,
    }


def extract_docx(path: Path) -> dict:
    try:
        from docx import Document
    except ImportError:
        print("Installing python-docx...", file=sys.stderr)
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
        from docx import Document

    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            tables.append(rows)

    full_text = "\n\n".join(paragraphs)
    if tables:
        full_text += "\n\n--- TABLES ---\n"
        for ti, table in enumerate(tables, 1):
            full_text += f"\nTable {ti}:\n"
            for row in table:
                full_text += " | ".join(row) + "\n"

    return {
        "source": path.name,
        "type": "docx",
        "paragraph_count": len(paragraphs),
        "paragraphs": paragraphs,
        "tables": tables,
        "full_text": full_text,
    }


def extract_html(path: Path) -> dict:
    html = path.read_text(encoding="utf-8", errors="replace")
    # Strip tags for plain text summary
    text = re.sub(r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return {
        "source": path.name,
        "type": "html",
        "full_text": text[:50000],
        "note": "Full wireframe HTML preserved separately; this is a text summary.",
    }


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    results: list[dict] = []

    for path in sorted(DOCS_DIR.iterdir()):
        if path.name == "extracted" or path.is_dir():
            continue
        suffix = path.suffix.lower()
        print(f"Extracting: {path.name}", file=sys.stderr)
        try:
            if suffix == ".pdf":
                data = extract_pdf(path)
            elif suffix == ".docx":
                data = extract_docx(path)
            elif suffix == ".html":
                data = extract_html(path)
            else:
                continue
            results.append(data)
            out_file = OUTPUT_DIR / f"{path.stem}.json"
            out_file.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")
            txt_file = OUTPUT_DIR / f"{path.stem}.txt"
            txt_file.write_text(data.get("full_text", ""), encoding="utf-8")
        except Exception as e:
            print(f"  ERROR: {e}", file=sys.stderr)
            results.append({"source": path.name, "error": str(e)})

    manifest = {
        "extracted_at": datetime.now(timezone.utc).isoformat(),
        "files": [{"source": r["source"], "type": r.get("type"), "error": r.get("error")} for r in results],
    }
    (OUTPUT_DIR / "manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    print(f"\nExtracted {len(results)} files to {OUTPUT_DIR}", file=sys.stderr)


if __name__ == "__main__":
    main()
